import os
from pdf2image import convert_from_bytes
from docx import Document
from .image_extractor import ImageExtractor
import io
import tempfile
import threading
import hashlib
from functools import lru_cache
from concurrent.futures import ThreadPoolExecutor
import logging
import fitz  # PyMuPDF
import time
from PIL import Image
import numpy as np
import re
from cachetools import LRUCache
from threading import Lock

logger = logging.getLogger(__name__)

class DocumentExtractor:
    def __init__(self):
        self.image_extractor = ImageExtractor()
        # Use LRUCache for better performance
        self.cache = LRUCache(maxsize=100)
        self.cache_lock = Lock()
        self.thread_pool = ThreadPoolExecutor(max_workers=8)
        # PDF extraction settings
        self.min_text_length = 50  # Minimum length to consider text extraction successful
        self.max_dpi = 300  # Maximum DPI to use for OCR
        self.min_dpi = 150  # Minimum DPI to use for OCR
        self.default_dpi = 200  # Default DPI for OCR
        self.ocr_threshold = 0.6  # Confidence threshold for OCR

    def compute_hash(self, file_bytes):
        """Compute hash of file bytes for caching"""
        return hashlib.sha256(file_bytes).hexdigest()

    def extract_text_from_pdf_direct(self, file_bytes):
        """Extract text directly from PDF without OCR"""
        try:
            # Load PDF from memory
            with fitz.open(stream=file_bytes, filetype="pdf") as pdf_document:
                total_pages = len(pdf_document)
                logger.info(f"PDF has {total_pages} pages")
                
                # Prepare futures for parallel processing
                futures = []
                for page_num in range(total_pages):
                    futures.append(self.thread_pool.submit(
                        self.process_pdf_page_direct, pdf_document, page_num
                    ))
                
                # Collect results
                page_texts = []
                for future in futures:
                    page_text = future.result()
                    if page_text:
                        page_texts.append(page_text)
                
                # Join all page texts
                full_text = "\n\n".join(page_texts)
                
                # Check if we got enough text
                if len(full_text.strip()) > self.min_text_length:
                    logger.info("Successfully extracted text directly from PDF")
                    return full_text, True
                else:
                    logger.info("Direct text extraction yielded insufficient text, will try OCR")
                    return full_text, False
                
        except Exception as e:
            logger.warning(f"Error in direct PDF text extraction: {str(e)}")
            return "", False

    def process_pdf_page_direct(self, pdf_document, page_num):
        """Process a single PDF page for text extraction with improved table handling"""
        try:
            page = pdf_document[page_num]
            
            # First, check if the page contains tables
            table_data = self.extract_tables_from_pdf_page(page)
            
            # Get regular text with enhanced options
            text = page.get_text("text", sort=True, flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE)
            
            # If text is too short, try alternative extraction methods
            if len(text.strip()) < 20:
                # Try extracting with different flags
                text = page.get_text("text", sort=True, flags=fitz.TEXT_DEHYPHENATE)
                
                # If still too short, try extracting as blocks
                if len(text.strip()) < 20:
                    blocks = page.get_text("blocks")
                    if blocks:
                        text = "\n".join([b[4] for b in blocks])
            
            # Clean the text
            text = self.clean_pdf_text(text)
            
            # Add formatted tables to the text
            if table_data:
                for table in table_data:
                    formatted_table = self.format_table(table)
                    text += "\n\n" + formatted_table
            
            return text
            
        except Exception as e:
            logger.warning(f"Error processing PDF page {page_num}: {str(e)}")
            return ""

    def extract_tables_from_pdf_page(self, page):
        """Extract tables from a PDF page using layout analysis"""
        try:
            # Get blocks which might represent tables
            blocks = page.get_text("dict")["blocks"]
            tables = []
            
            # Look for potential table structures
            for block in blocks:
                # Check if block contains lines (potential table)
                if "lines" in block:
                    lines = block["lines"]
                    
                    # Skip blocks with too few lines
                    if len(lines) < 2:
                        continue
                    
                    # Check if lines have consistent spans (potential table rows)
                    row_data = []
                    for line in lines:
                        if "spans" in line:
                            # Extract text from spans
                            cell_data = []
                            for span in line["spans"]:
                                if "text" in span and span["text"].strip():
                                    cell_data.append(span["text"].strip())
                            
                            if cell_data:
                                row_data.append(cell_data)
                    
                    # If we have multiple rows with data, consider it a table
                    if len(row_data) >= 2:
                        # Normalize the table (ensure all rows have same number of columns)
                        max_cols = max(len(row) for row in row_data)
                        if max_cols >= 2:  # Only consider as table if at least 2 columns
                            normalized_rows = []
                            for row in row_data:
                                if len(row) >= max_cols * 0.5:  # Row must have at least half the columns to be valid
                                    normalized_rows.append(row + [''] * (max_cols - len(row)))
                            
                            if normalized_rows:
                                tables.append(normalized_rows)
            
            return tables
            
        except Exception as e:
            logger.warning(f"Error extracting tables from PDF page: {str(e)}")
            return []

    def clean_pdf_text(self, text):
        """Clean text extracted from PDF"""
        if not text:
            return ""
            
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines while preserving paragraph breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove non-printable characters
        text = ''.join(char for char in text if char.isprintable() or char == '\n')
        
        return text.strip()

    def format_table(self, table_data):
        """
        Format a table into readable text.
        
        Args:
            table_data: List of lists representing table rows and cells
            
        Returns:
            Formatted text representation of the table
        """
        if not table_data or not table_data[0]:
            return ""
            
        # Calculate column widths
        col_count = max(len(row) for row in table_data)
        col_widths = [0] * col_count
        
        # Standardize all rows to have the same number of columns
        for i in range(len(table_data)):
            # Extend rows that have fewer cells than the maximum
            if len(table_data[i]) < col_count:
                table_data[i].extend([''] * (col_count - len(table_data[i])))
        
        # Determine the maximum width needed for each column
        for row in table_data:
            for i, cell in enumerate(row):
                col_widths[i] = max(col_widths[i], len(str(cell).strip()))
        
        # Format the table with proper spacing
        result = []
        
        # Add a header separator after the first row
        header_sep = "+"
        for width in col_widths:
            header_sep += "-" * (width + 2) + "+"
        
        # Format each row
        for i, row in enumerate(table_data):
            formatted_row = "| "
            for j, cell in enumerate(row):
                cell_text = str(cell).strip()
                formatted_row += cell_text.ljust(col_widths[j]) + " | "
            result.append(formatted_row)
            
            # Add separator after header
            if i == 0:
                result.append(header_sep)
        
        # Add a newline before and after the table
        return "\n" + "\n".join(result) + "\n"

    def extract_from_docx(self, file_bytes):
        """Extract text from DOCX files with improved table handling"""
        try:
            # Check cache first
            file_hash = self.compute_hash(file_bytes)
            with self.cache_lock:
                if file_hash in self.cache:
                    logger.info("Using cached DOCX result")
                    return self.cache[file_hash]
            
            # Save bytes to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name

            # Open and extract text from document
            doc = Document(tmp_path)
            text_blocks = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():  # Only add non-empty paragraphs
                    text_blocks.append(para.text)
            
            # Extract text from tables with improved formatting
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        # Get all text from the cell, including from paragraphs
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_data.append(cell_text)
                        else:
                            row_data.append("")  # Empty cell
                    table_data.append(row_data)
                
                # Format the table and add to text blocks
                if table_data:
                    formatted_table = self.format_table(table_data)
                    text_blocks.append(formatted_table)
            
            # Clean up temporary file
            os.unlink(tmp_path)
            
            result = '\n\n'.join(text_blocks)
            
            # Cache the result
            with self.cache_lock:
                self.cache[file_hash] = result
            
            return result
            
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")

    def estimate_optimal_dpi(self, pdf_document):
        """Estimate optimal DPI based on document complexity"""
        try:
            # Sample a few pages to determine complexity
            num_pages = len(pdf_document)
            sample_pages = min(3, num_pages)
            
            total_chars = 0
            total_images = 0
            
            for i in range(min(sample_pages, num_pages)):
                page = pdf_document[i]
                
                # Count text characters
                text = page.get_text("text")
                total_chars += len(text)
                
                # Count images
                image_list = page.get_images()
                total_images += len(image_list)
            
            # Calculate average complexity
            avg_chars = total_chars / sample_pages
            avg_images = total_images / sample_pages
            
            # Determine DPI based on complexity
            if avg_chars > 1000 or avg_images > 5:
                # Complex document with lots of text or images
                return self.max_dpi
            elif avg_chars < 200 and avg_images < 2:
                # Simple document with little text and few images
                return self.min_dpi
            else:
                # Medium complexity
                return self.default_dpi
                
        except Exception as e:
            logger.warning(f"Error estimating optimal DPI: {str(e)}")
            return self.default_dpi

    def extract_from_pdf(self, file_bytes):
        """Extract text from PDF using both direct extraction and OCR with optimizations"""
        try:
            start_time = time.time()
            
            # Check cache first
            file_hash = self.compute_hash(file_bytes)
            with self.cache_lock:
                if file_hash in self.cache:
                    logger.info("Using cached PDF result")
                    return self.cache[file_hash]
            
            # Step 1: Try direct text extraction first (much faster)
            direct_text, direct_success = self.extract_text_from_pdf_direct(file_bytes)
            
            # If direct extraction was successful, return the result
            if direct_success:
                # Cache the result
                with self.cache_lock:
                    self.cache[file_hash] = direct_text
                
                logger.info(f"PDF extraction completed in {time.time() - start_time:.2f}s using direct extraction")
                return direct_text
            
            # Step 2: If direct extraction failed, use OCR with optimized parameters
            logger.info("Direct extraction insufficient, falling back to OCR")
            
            # Determine optimal DPI based on document complexity
            with fitz.open(stream=file_bytes, filetype="pdf") as pdf_document:
                optimal_dpi = self.estimate_optimal_dpi(pdf_document)
                logger.info(f"Using optimal DPI: {optimal_dpi}")
            
            # Convert PDF to images with optimized parameters
            images = convert_from_bytes(
                file_bytes,
                dpi=optimal_dpi,
                thread_count=8,  # Use more threads for conversion
                use_pdftocairo=True,  # Use pdftocairo which is faster
                grayscale=True,  # Convert to grayscale for better OCR
                size=(None, 2000)  # Limit height to 2000px for speed while maintaining quality
            )
            
            # Extract text from each page in parallel with improved batching
            # Process images in batches to avoid memory issues
            batch_size = 4
            all_text = []
            
            for i in range(0, len(images), batch_size):
                batch = images[i:i+batch_size]
                
                # Process batch in parallel
                futures = []
                for image in batch:
                    futures.append(self.thread_pool.submit(
                        self.process_pdf_page_ocr, image, i
                    ))
                
                # Collect batch results
                for future in futures:
                    try:
                        page_text = future.result()
                        if page_text:
                            all_text.append(page_text)
                    except Exception as e:
                        logger.warning(f"Error extracting text from PDF page: {str(e)}")
            
            # Combine text from all pages
            if direct_text and len(direct_text) > self.min_text_length:
                # If we have some direct text, combine it with OCR results
                result = direct_text + "\n\n" + "\n\n".join(all_text)
            else:
                # Otherwise just use OCR results
                result = "\n\n".join(all_text)
            
            # Post-process the combined text
            result = self.post_process_pdf_text(result)
            
            # Cache the result
            with self.cache_lock:
                self.cache[file_hash] = result
            
            logger.info(f"PDF extraction completed in {time.time() - start_time:.2f}s using OCR")
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise Exception(f"Error extracting text from PDF: {str(e)}")

    def process_pdf_page_ocr(self, image, page_num):
        """Process a single PDF page with OCR"""
        try:
            # Enhanced preprocessing for PDF images
            enhanced_image = self.enhance_pdf_image(image)
            
            # Extract text using the image extractor
            text = self.image_extractor.extract_text_from_pdf_image(enhanced_image)
            
            # Add page number for reference
            return f"Page {page_num + 1}:\n{text}"
            
        except Exception as e:
            logger.warning(f"Error processing PDF page {page_num} with OCR: {str(e)}")
            return ""

    def enhance_pdf_image(self, image):
        """Apply PDF-specific enhancements to improve OCR quality"""
        try:
            # Convert PIL image to numpy array
            img_array = np.array(image)
            
            # Check if image is already grayscale
            if len(img_array.shape) == 3 and img_array.shape[2] == 3:
                # Convert to grayscale
                gray = np.dot(img_array[...,:3], [0.2989, 0.5870, 0.1140])
                gray = gray.astype(np.uint8)
            else:
                gray = img_array
            
            # Create PIL image from array
            enhanced = Image.fromarray(gray)
            
            # Increase contrast slightly
            enhancer = Image.ImageEnhance.Contrast(enhanced)
            enhanced = enhancer.enhance(1.5)
            
            # Increase sharpness
            enhancer = Image.ImageEnhance.Sharpness(enhanced)
            enhanced = enhancer.enhance(1.5)
            
            return enhanced
            
        except Exception as e:
            logger.warning(f"Error enhancing PDF image: {str(e)}")
            return image  # Return original if enhancement fails

    def post_process_pdf_text(self, text):
        """Apply post-processing to improve extracted PDF text with better table handling"""
        if not text:
            return ""
            
        # Remove duplicate lines that often occur in PDFs
        lines = text.split('\n')
        unique_lines = []
        prev_line = ""
        
        # Flag to identify table sections
        in_table = False
        
        for line in lines:
            # Check if we're entering or leaving a table section
            if line.startswith('+') and ('-' in line):
                in_table = True
                unique_lines.append(line)
                prev_line = line
                continue
            
            # If we're in a table, preserve all lines including empty ones
            if in_table:
                if not line.strip() and not prev_line.strip():
                    in_table = False  # End of table
                else:
                    unique_lines.append(line)
                    prev_line = line
                continue
            
            # Skip empty lines
            if not line.strip():
                if prev_line:  # Only add newline if we have content
                    unique_lines.append("")
                prev_line = ""
                continue
                
            # Skip duplicate lines
            if line.strip() == prev_line.strip():
                continue
                
            unique_lines.append(line)
            prev_line = line
        
        # Join lines back together
        text = '\n'.join(unique_lines)
        
        # Fix common OCR issues
        text = re.sub(r'([a-z])- ([a-z])', r'\1\2', text)  # Fix hyphenated words
        
        # Don't normalize whitespace within table sections
        processed_lines = []
        in_table = False
        
        for line in text.split('\n'):
            if line.startswith('+') and ('-' in line):
                in_table = True
                processed_lines.append(line)
                continue
                
            if in_table:
                if not line.strip():
                    in_table = False
                    processed_lines.append(line)
                else:
                    processed_lines.append(line)
                continue
                
            # For non-table text, normalize whitespace
            processed_line = re.sub(r'\s+', ' ', line)
            processed_lines.append(processed_line)
        
        text = '\n'.join(processed_lines)
        
        # Normalize newlines (but not in tables)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()

    def extract_from_txt(self, file_bytes):
        """Extract text from TXT files"""
        try:
            # Check cache first
            file_hash = self.compute_hash(file_bytes)
            with self.cache_lock:
                if file_hash in self.cache:
                    logger.info("Using cached TXT result")
                    return self.cache[file_hash]
            
            # Decode bytes with different encodings
            encodings = self.get_encoding_list()
            text = None
            
            for encoding in encodings:
                try:
                    text = file_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise Exception("Could not decode text file with supported encodings")
            
            # Cache the result
            with self.cache_lock:
                self.cache[file_hash] = text
            
            return text
            
        except Exception as e:
            raise Exception(f"Error extracting text from TXT: {str(e)}")

    def get_encoding_list(self):
        """Return list of encodings to try, cached to avoid recreation"""
        return ['utf-8', 'latin-1', 'ascii', 'utf-16', 'windows-1252']

    def extract_text_from_document(self, file_bytes, file_extension):
        """Main method to extract text based on file type"""
        extractors = {
            '.pdf': self.extract_from_pdf,
            '.docx': self.extract_from_docx,
            '.txt': self.extract_from_txt
        }
        
        if file_extension not in extractors:
            raise ValueError(f"Unsupported document type: {file_extension}")
        
        return extractors[file_extension](file_bytes)