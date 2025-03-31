from fastapi import FastAPI, UploadFile, HTTPException
from pydantic import BaseModel, Field, validator
import requests
import os
import aiohttp
from dotenv import load_dotenv
import uuid
import google.generativeai as genai  
import json
import traceback
from fastapi.responses import JSONResponse, StreamingResponse
from datetime import datetime, timedelta
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from io import BytesIO
import zipfile
import re
import logging
import time
logger = logging.getLogger(__name__)
load_dotenv()

app = FastAPI()

# JIRA Configuration
JIRA_URL = os.getenv("JIRA_URL")
JIRA_USER = os.getenv("JIRA_USER")
JIRA_API_TOKEN = os.getenv("JIRA_API_TOKEN")
PROJECT_KEY = os.getenv("JIRA_PROJECT_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

genai.configure(api_key=GEMINI_API_KEY)

class JiraIssue(BaseModel):
    summary: str
    description: str
    issue_type: str = Field(default="Story", pattern=r"^(Story|Task|Bug)$")
    priority: str = Field(default="Medium", pattern=r"^(High|Medium|Low)$")

    @validator('priority', pre=True)
    def set_default_priority(cls, value):
        return value or "Medium"  

class EpicRequest(BaseModel):
    epic_summary: str
    epic_description: str
    story_keys: list[str] = []  

class PriorityRequest(BaseModel):
    strategy: str = "business_value"  

class TimelineRequest(BaseModel):
    start_date: datetime
    sprint_duration: int = Field(default=14, description="Sprint duration in days")
    team_capacity: int = Field(default=80, description="Team capacity in story points per sprint")

class DocumentRequest(BaseModel):
    format: str = Field(default="word", pattern="^(word|excel|both)$")
    include_metadata: bool = Field(default=True)
    template_type: str = Field(default="detailed", pattern="^(detailed|simple)$")

def convert_to_adf(text: str) -> dict:
    """Convert plain text with line breaks to ADF"""
    content = []
    for paragraph in text.split('\n'):
        if paragraph.strip():
            content.append({
                "type": "paragraph",
                "content": [
                    {
                        "type": "text",
                        "text": paragraph.strip()
                    }
                ]
            })
    return {
        "version": 1,
        "type": "doc",
        "content": content
    }

def create_jira_issue(issue: JiraIssue):
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json"
    }

    payload = {
        "fields": {
            "project": {"key": "ARE"},  
            "summary": issue.summary,
            "description": convert_to_adf(issue.description),
            "issuetype": {"name": issue.issue_type}
        }
    }

    if issue.priority:
        payload["fields"]["priority"] = {"name": issue.priority}

    response = requests.post(
        f"{JIRA_URL}/rest/api/3/issue",
        auth=(JIRA_USER, JIRA_API_TOKEN),
        headers=headers,
        json=payload
    )

    if response.status_code != 201:
        raise HTTPException(status_code=response.status_code, detail=response.text)
    
    return response.json()

def create_epic(epic_data: EpicRequest):
    """Create an epic in JIRA Next-gen project"""
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json"
    }

    payload = {
        "fields": {
            "project": {"key": "ARE"},
            "summary": epic_data.epic_summary,
            "description": convert_to_adf(epic_data.epic_description),
            "issuetype": {"name": "Epic"}
        }
    }

    response = requests.post(
        f"{JIRA_URL}/rest/api/3/issue",
        auth=(JIRA_USER, JIRA_API_TOKEN),
        headers=headers,
        json=payload
    )

    if response.status_code != 201:
        raise HTTPException(status_code=response.status_code, detail=response.text)
    
    return response.json()

def link_issues(epic_key: str, story_key: str):
    """Link a story to an epic in a Next-gen project"""
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json"
    }

    payload = {
        "fields": {
            "parent": {
                "key": epic_key
            }
        }
    }

    response = requests.put(
        f"{JIRA_URL}/rest/api/3/issue/{story_key}",
        auth=(JIRA_USER, JIRA_API_TOKEN),
        headers=headers,
        json=payload
    )

    if response.status_code not in [200, 204]:
        raise HTTPException(status_code=response.status, detail=response.text)

@app.post("/upload-srs/")
async def upload_srs(file: UploadFile):
    """Process SRS document and create JIRA issues"""
    try:
        content = await file.read()
        print("\n" + "="*40 + " ORIGINAL SRS CONTENT " + "="*40)
        print(content.decode())
        print("="*95 + "\n")
        
        requirements = process_srs(content.decode())
        
        created_issues = []
        for req in requirements:
            try:
                print("\n" + "="*40 + " CREATING ISSUE " + "="*40)
                print("Request payload:", req)
                
                issue = JiraIssue(
                    summary=req["summary"],
                    description=req["description"],
                    priority=req.get("priority", "Medium")
                )
                result = create_jira_issue(issue)
                created_issues.append(result)
                
                print("Created issue:", result)
                print("="*95 + "\n")
                
            except Exception as e:
                print(f"Failed to create issue: {str(e)}")
                continue
                
        return {"status": "success", "created_issues": created_issues}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

def identify_domain(content: str) -> str:
    """Use Gemini to identify the domain from the SRS content"""
    model = genai.GenerativeModel("gemini-1.5-pro")
    domain_prompt = f"""Analyze this software requirements document and identify the specific domain (e.g., Healthcare, Banking, Education, etc.).
    Return ONLY the domain name, nothing else.
    
    Document:
    {content[:2000]}  # First 2000 chars should be enough for domain identification
    """
    try:
        response = model.generate_content(domain_prompt)
        domain = response.text.strip()
        return domain
    except Exception as e:
        logger.error(f"Error identifying domain: {e}")
        return "General Software System"

def process_srs(file_content: str):
    """Process structured SRS document and extract requirements"""
    requirements = []
    
    # First identify the domain
    domain = identify_domain(file_content)
    logger.info(f"Identified domain: {domain}")
    
    sections = file_content.split("##")
    req_pattern = r'\[((?:FR|NFR|SR)-\d+)\]:\s*(.*?)\s*\[(High|Medium|Low)\]'
    
    raw_requirements = []
    for section in sections:
        matches = re.finditer(req_pattern, section, re.MULTILINE)
        for match in matches:
            req_id, description, priority = match.groups()
            raw_requirements.append({
                "id": req_id,
                "description": description,
                "priority": priority,
                "type": req_id.split('-')[0],
                "domain": domain
            })
    
    if not raw_requirements:
        return []

    # Process requirements in smaller batches
    BATCH_SIZE = 3
    for i in range(0, len(raw_requirements), BATCH_SIZE):
        batch = raw_requirements[i:i + BATCH_SIZE]
        
        prompt = f"""As an expert in requirements engineering for {domain} systems, enhance these software requirements with clear, professional descriptions and titles.
        For each requirement, provide:
        1. title: A concise, action-oriented title that clearly states the requirement's purpose (max 70 chars)
        2. description: A well-structured description that includes:
           - Clear objective
           - User benefit/business value
           - Technical context
           - Implementation considerations
           - Success criteria
           - Any constraints or dependencies
        
        Format each requirement in a JSON array with this structure:
        {{
            "id": "original_id",
            "title": "enhanced_title",
            "description": "enhanced_description"
        }}

        Original requirements to enhance:
        {json.dumps(batch, indent=2)}
        """
        
        try:
            model = genai.GenerativeModel('gemini-1.5-pro')
            response = model.generate_content(prompt, timeout=30)  # 30 second timeout
            
            # Parse the response
            enhanced = parse_ai_response(response.text)
            if isinstance(enhanced, list):
                requirements.extend(enhanced)
            else:
                print(f"Error parsing batch {i}-{i+BATCH_SIZE}: Invalid response format")
                # Use original requirements as fallback
                for req in batch:
                    requirements.append({
                        "id": req["id"],
                        "title": f"{req['type']} - {req['description'][:50]}...",
                        "description": req["description"]
                    })
        except Exception as e:
            print(f"Error processing batch {i}-{i+BATCH_SIZE}: {str(e)}")
            # Use original requirements as fallback
            for req in batch:
                requirements.append({
                    "id": req["id"],
                    "title": f"{req['type']} - {req['description'][:50]}...",
                    "description": req["description"]
                })
        
        # Add a small delay between batches
        time.sleep(1)
    
    return requirements

@app.post("/prioritize-backlog")
async def prioritize_backlog(strategy: PriorityRequest):
    """AI-driven backlog prioritization with multiple strategies"""
    try:
        print("\n=== STARTING PRIORITIZATION ===")
        
        print("Verifying JIRA configuration...")
        if not all([JIRA_URL, JIRA_USER, JIRA_API_TOKEN, "ARE"]):
            raise ValueError("Missing JIRA configuration in environment variables")
            
        print(f"Fetching backlog for project ARE...")
        async with aiohttp.ClientSession() as session:
            async with session.get(
                f"{JIRA_URL}/rest/api/3/search",
                auth=aiohttp.BasicAuth(JIRA_USER, JIRA_API_TOKEN),
                params={"jql": "project=ARE"}
            ) as response:
                if response.status != 200:
                    text = await response.text()
                    raise HTTPException(status_code=response.status, detail=text)
                
                data = await response.json()
                issues = data.get("issues", [])
        
        print(f"Found {len(issues)} backlog items")
        
        print("Preparing data for AI analysis...")
        issue_data = [{
            "key": issue["key"],
            "summary": issue["fields"]["summary"],
            "description": issue["fields"].get("description", ""),
            "current_priority": issue["fields"]["priority"].get("name", "Medium")
        } for issue in issues]
        
        prompt = f"""Analyze these JIRA issues and prioritize them based on {strategy.strategy}.
Return ONLY a JSON array with this exact structure:
[
    {{
        "issue_key": "PROJ-123", 
        "new_priority": "High",
        "rationale": "1-2 sentence explanation"
    }},
    ...
]
Do NOT include any other text or formatting. Issues:
{json.dumps(issue_data, indent=2)}"""
        
        print("\n=== AI PROMPT ===")
        print(prompt[:1000] + "..." if len(prompt) > 1000 else prompt)
        
        print("Querying Gemini AI...")
        model = genai.GenerativeModel("gemini-1.5-pro")
        response = model.generate_content(prompt)
        raw_response = response.text.replace('```json', '').replace('```', '').strip()
        
        print("\n=== AI RAW RESPONSE ===")
        print(raw_response[:1000] + "..." if len(raw_response) > 1000 else raw_response)
        
        print("Parsing AI response...")
        try:
            prioritized_list = parse_ai_response(raw_response)
            if not isinstance(prioritized_list, list):
                raise ValueError("AI response is not a list")
        except Exception as e:
            print(f"JSON Parsing Error: {str(e)}")
            raise
        
        valid_priorities = ["Highest", "High", "Medium", "Low", "Lowest"]
        for item in prioritized_list:
            if item["new_priority"] not in valid_priorities:
                raise ValueError(f"Invalid priority {item['new_priority']} for {item['issue_key']}")
        
        print("Updating JIRA priorities...")
        success_count = 0
        for item in prioritized_list:
            try:
                async with aiohttp.ClientSession() as session:
                    async with session.put(
                        f"{JIRA_URL}/rest/api/3/issue/{item['issue_key']}",
                        auth=aiohttp.BasicAuth(JIRA_USER, JIRA_API_TOKEN),
                        json={"fields": {"priority": {"name": item["new_priority"]}}}
                    ) as response:
                        if response.status != 204:
                            text = await response.text()
                            print(f"Failed to update {item['issue_key']}: {text}")
                        else:
                            success_count += 1
                    
            except Exception as e:
                print(f"Error updating {item['issue_key']}: {str(e)}")
        
        return {
            "status": "partial" if success_count < len(prioritized_list) else "success",
            "updated": success_count,
            "failed": len(prioritized_list) - success_count
        }

    except Exception as e:
        print("\n=== ERROR DETAILS ===")
        print(f"Type: {type(e).__name__}")
        print(f"Message: {str(e)}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail={
            "error": str(e),
            "type": type(e).__name__
        })

@app.post("/create-epic/")
async def create_epic_endpoint(epic_request: EpicRequest):
    """Create an epic and link stories to it"""
    try:
        epic_result = create_epic(epic_request)
        epic_key = epic_result["key"]
        
        linked_stories = []
        for story_key in epic_request.story_keys:
            try:
                link_issues(epic_key, story_key)
                linked_stories.append(story_key)
            except Exception as e:
                print(f"Failed to link story {story_key}: {str(e)}")
        
        return {
            "status": "success",
            "epic_key": epic_key,
            "linked_stories": linked_stories,
            "failed_stories": [key for key in epic_request.story_keys if key not in linked_stories]
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to create epic or link stories: {str(e)}"
        )

@app.post("/automate-workflow/")
async def automate_workflow():
    """Automate task assignment and status updates"""
    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(
                f"{JIRA_URL}/rest/api/3/search",
                auth=aiohttp.BasicAuth(JIRA_USER, JIRA_API_TOKEN),
                params={"jql": f"project={PROJECT_KEY} AND status='To Do'"}
            ) as response:
                if response.status != 200:
                    text = await response.text()
                    raise HTTPException(status_code=response.status, detail=text)
                
                data = await response.json()
                issues = data.get("issues", [])
        
        for issue in issues:
            issue_key = issue["key"]
            transition_payload = {"transition": {"id": "21"}}  
            
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    f"{JIRA_URL}/rest/api/3/issue/{issue_key}/transitions",
                    auth=aiohttp.BasicAuth(JIRA_USER, JIRA_API_TOKEN),
                    json=transition_payload
                ) as response:
                    if response.status != 204:
                        text = await response.text()
                        raise HTTPException(status_code=response.status, detail=text)
        
        return {"status": "success", "message": "Workflow automated successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

async def get_jira_issues(jql: str = "issuetype = Story"):
    """Fetch issues from JIRA based on JQL"""
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json"
    }
    
    params = {
        "jql": jql,
        "maxResults": 100  
    }
    
    async with aiohttp.ClientSession() as session:
        async with session.get(
            f"{JIRA_URL}/rest/api/3/search",
            auth=aiohttp.BasicAuth(JIRA_USER, JIRA_API_TOKEN),
            headers=headers,
            params=params
        ) as response:
            if response.status != 200:
                text = await response.text()
                raise HTTPException(status_code=response.status, detail=text)
            
            data = await response.json()
            return data.get("issues", [])

def suggest_epics(stories: list) -> list:
    """Use Gemini AI to suggest epic groupings for stories"""
    stories_text = "\n".join([
        f"Story {story['key']}:\n"
        f"Title: {story['fields']['summary']}\n"
        f"Description: {story['fields'].get('description', 'No description')}\n"
        f"Priority: {story['fields'].get('priority', {}).get('name', 'Not set')}\n"
        f"Labels: {', '.join(story['fields'].get('labels', []))}\n"
        for story in stories
    ])
    
    prompt = f"""As an expert in Agile project management and JIRA organization, analyze these user stories and suggest optimal epic groupings.

    Consider these key factors when grouping stories into epics:
    1. Business Value Stream:
       - Group stories that contribute to the same business objective
       - Consider end-user value and customer journey
       - Identify common business processes or workflows
    
    2. Technical Dependencies:
       - Identify stories that share technical components
       - Consider architectural layers (UI, backend, database)
       - Group stories with similar technical requirements
    
    3. Feature Completeness:
       - Ensure each epic represents a complete, deliverable feature
       - Consider MVP (Minimum Viable Product) requirements
       - Include all necessary stories for end-to-end functionality
    
    4. Timeline and Priority:
       - Group stories that should be implemented in the same timeframe
       - Consider dependencies and sequential implementation needs
       - Balance epic sizes for manageable delivery
    
    For each suggested epic, provide:
    1. epic_summary: A clear, concise title (max 70 chars)
    2. epic_description: A detailed description including:
       - Epic's main objective
       - Key deliverables
       - Success criteria
       - Technical considerations
       - Dependencies on other epics
    3. story_keys: List of story keys to include
    4. rationale: Detailed explanation including:
       - Why these stories belong together
       - Business value alignment
       - Technical synergies
       - Implementation considerations
    5. estimated_duration: Rough time estimate (in sprints)
    6. suggested_priority: High/Medium/Low based on business impact
    
    Return the response as a JSON array where each epic follows this structure:
    {{
        "epic_summary": "string",
        "epic_description": "string",
        "story_keys": ["key1", "key2"],
        "rationale": "string",
        "estimated_duration": "number",
        "suggested_priority": "string"
    }}

    Here are the stories to analyze:
    {stories_text}
    """
    
    model = genai.GenerativeModel("gemini-1.5-pro")
    response = model.generate_content(prompt)
    
    try:
        return parse_ai_response(response.text)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to parse AI suggestions: {str(e)}"
        )

@app.post("/suggest-epics/")
async def suggest_epics_endpoint(jql: str = "issuetype = Story"):
    """Get AI suggestions for grouping stories into epics"""
    try:
        # Fetch stories from JIRA
        stories = await get_jira_issues(jql)
        
        if not stories:
            return {"message": "No stories found matching the criteria"}
        
        # Get AI suggestions
        suggestions = suggest_epics(stories)
        
        return {
            "status": "success",
            "epic_suggestions": suggestions,
            "total_stories_analyzed": len(stories)
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate epic suggestions: {str(e)}"
        )

@app.post("/create-suggested-epics/")
async def create_suggested_epics(jql: str = "issuetype = Story"):
    """Get AI suggestions and automatically create the suggested epics"""
    try:
        suggestions = await suggest_epics_endpoint(jql)
        
        if "epic_suggestions" not in suggestions:
            return suggestions
        
        created_epics = []
        for epic in suggestions["epic_suggestions"]:
            epic_request = EpicRequest(
                epic_summary=epic["epic_summary"],
                epic_description=f"{epic['epic_description']}\n\nRationale: {epic['rationale']}",
                story_keys=epic["story_keys"]
            )
            
            result = await create_epic_endpoint(epic_request)
            created_epics.append(result)
        
        return {
            "status": "success",
            "created_epics": created_epics,
            "total_epics_created": len(created_epics)
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to create suggested epics: {str(e)}"
        )

async def generate_epic_timeline(epics: list, timeline_req: TimelineRequest):
    """Generate timeline for epics using GenAI"""
    epics_data = []
    for epic in epics:
        epic_data = {
            "key": epic["key"],
            "summary": epic["fields"]["summary"],
            "description": epic["fields"].get("description", ""),
            "priority": epic["fields"].get("priority", {}).get("name", "Medium"),
            "stories": []
        }
        
        jql = f'parent = {epic["key"]}'
        stories = await get_jira_issues(jql)
        epic_data["stories"] = [{
            "key": story["key"],
            "summary": story["fields"]["summary"],
            "story_points": story["fields"].get("customfield_10014", 0)  
        } for story in stories]
        
        epics_data.append(epic_data)

    prompt = f"""As an expert in Agile project planning, create an optimal timeline for these epics starting from {timeline_req.start_date.strftime('%Y-%m-%d')}.

    Project Context:
    - Sprint duration: {timeline_req.sprint_duration} days
    - Team capacity: {timeline_req.team_capacity} story points per sprint
    
    For each epic, analyze:
    1. Scope and complexity
    2. Dependencies between stories
    3. Priority and business value
    4. Required team capacity
    5. Technical dependencies
    
    Epics and their stories:
    {json.dumps(epics_data, indent=2)}
    
    Return a JSON array of epic schedules with:
    {{
        "epic_key": "string",
        "start_date": "YYYY-MM-DD",
        "end_date": "YYYY-MM-DD",
        "estimated_sprints": number,
        "story_point_total": number,
        "parallel_epics": ["epic_keys"],
        "scheduling_rationale": "string"
    }}
    
    Consider:
    - Story point distribution across sprints
    - Dependencies between epics
    - Parallel execution possibilities
    - Team capacity constraints
    - Priority order
    """
    
    model = genai.GenerativeModel("gemini-1.5-pro")
    response = model.generate_content(prompt)
    
    try:
        timeline = parse_ai_response(response.text)
        
        for schedule in timeline:
            epic_key = schedule["epic_key"]
            
            update_payload = {
                "fields": {
                    "customfield_10015": schedule["start_date"],  
                    "duedate": schedule["end_date"],  
                    "description": {
                        "version": 1,
                        "type": "doc",
                        "content": [
                            {
                                "type": "paragraph",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": f"Estimated Duration: {schedule['estimated_sprints']} sprints\n"
                                    }
                                ]
                            },
                            {
                                "type": "paragraph",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": f"Story Points: {schedule['story_point_total']}\n"
                                    }
                                ]
                            },
                            {
                                "type": "paragraph",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": f"Parallel Epics: {', '.join(schedule['parallel_epics'])}\n"
                                    }
                                ]
                            },
                            {
                                "type": "paragraph",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": f"\nScheduling Rationale:\n{schedule['scheduling_rationale']}"
                                    }
                                ]
                            }
                        ]
                    }
                }
            }
            
            headers = {
                "Accept": "application/json",
                "Content-Type": "application/json"
            }
            
            async with aiohttp.ClientSession() as session:
                async with session.put(
                    f"{JIRA_URL}/rest/api/3/issue/{epic_key}",
                    auth=aiohttp.BasicAuth(JIRA_USER, JIRA_API_TOKEN),
                    headers=headers,
                    json=update_payload
                ) as response:
                    if response.status != 204:
                        text = await response.text()
                        print(f"Failed to update epic {epic_key}: {text}")
        
        return timeline
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate timeline: {str(e)}"
        )

@app.post("/generate-timeline")
async def generate_timeline_endpoint(timeline_req: TimelineRequest):
    """Generate and apply timeline for all epics"""
    try:
        epics = await get_jira_issues("issuetype = Epic")
        
        timeline = await generate_epic_timeline(epics, timeline_req)
        
        return {
            "status": "success",
            "timeline": timeline,
            "start_date": timeline_req.start_date,
            "total_epics": len(timeline)
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate timeline: {str(e)}"
        )

async def generate_word_document(requirements: list) -> BytesIO:
    """Generate a formatted Word document from requirements"""
    doc = Document()
    
    title = doc.add_heading('Software Requirements Specification', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Generated Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Total Requirements: {len(requirements)}')
    
    doc.add_paragraph('Table of Contents', style='Heading 1')
    doc.add_paragraph('(Update table of contents after opening the document)')
    
    req_types = {'FR': 'Functional Requirements',
                'NFR': 'Non-Functional Requirements',
                'SR': 'Security Requirements'}
    
    for req_type, section_title in req_types.items():
        type_reqs = [r for r in requirements if r['type'].startswith(req_type)]
        if type_reqs:
            doc.add_heading(section_title, level=1)
            
            for req in type_reqs:
                heading = doc.add_heading(level=2)
                heading.add_run(f"{req['summary']}")
                
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                table.autofit = True
                
                table.columns[0].width = Inches(2)
                table.columns[1].width = Inches(4)
                
                row_cells = table.rows[0].cells
                row_cells[0].text = 'Requirement ID'
                row_cells[1].text = req.get('id', 'N/A')
                
                details = [
                    ('Priority', req.get('priority', 'N/A')),
                    ('Type', req.get('type', 'N/A')),
                    ('Description', req.get('description', 'N/A')),
                ]
                
                for label, value in details:
                    row_cells = table.add_row().cells
                    row_cells[0].text = label
                    row_cells[1].text = str(value)
                
                doc.add_paragraph()  
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

async def parse_description_fields(description: str | dict) -> dict:
    """Parse structured fields from description"""
    fields = {
        'Objective': '',
        'User Benefit': '',
        'Technical Context': '',
        'Implementation Considerations': '',
        'Success Criteria': '',
        'Constraints': ''
    }
    
    if isinstance(description, dict):
        description = convert_adf_to_text(description)
    elif description is None:
        description = ''
    
    if description:
        for field in fields.keys():
            pattern = rf"\*\*{field}:\*\*(.*?)(?=\*\*\w+:\*\*|$)"
            match = re.search(pattern, description, re.DOTALL)
            if match:
                fields[field] = match.group(1).strip()
    
    return fields

def convert_adf_to_text(adf: dict) -> str:
    """Convert Atlassian Document Format to plain text"""
    text = []
    
    if not adf or 'content' not in adf:
        return ''
    
    for item in adf.get('content', []):
        if item.get('type') == 'paragraph':
            paragraph_text = []
            for content in item.get('content', []):
                if content.get('type') == 'text':
                    paragraph_text.append(content.get('text', ''))
            text.append(''.join(paragraph_text))
    
    return '\n'.join(text)

def extract_original_requirement(description: str | dict) -> str:
    """Extract the original requirement value from the description"""
    if isinstance(description, dict):
        description = convert_adf_to_text(description)
    elif description is None:
        return ''
    
    pattern = r'\[((?:FR|NFR|SR)-\d+)\]:\s*(.*?)\s*\[(High|Medium|Low)\]'
    match = re.search(pattern, description)
    if match:
        return match.group(2).strip()
    return description  

async def generate_excel_sheet(requirements: list) -> BytesIO:
    """Generate an Excel sheet from requirements for JIRA import"""

    excel_data = []
    for req in requirements:

        epic_link = ''
        if req['type'] == 'FR':
            epic_link = 'ARE-123'  
        elif req['type'] == 'NFR':
            epic_link = 'ARE-124'  
        elif req['type'] == 'SR':
            epic_link = 'ARE-125' 
        
        original_req = extract_original_requirement(req.get('description', ''))
        
        components = {
            'FR': 'Functional',
            'NFR': 'Non-Functional',
            'SR': 'Security'
        }.get(req['type'], '')
        
        excel_data.append({
            'Summary': req['summary'],
            'Description': original_req,
            'Priority': req['priority'],
            'Components': components,
            'Original Requirement ID': req.get('id', ''),
            'Epic Link': epic_link,
            'Is_Epic': req.get('fields', {}).get('issuetype', {}).get('name') == 'Epic'  
        })
    
    df = pd.DataFrame(excel_data)
    is_epic = df.pop('Is_Epic')  
    
    excel_io = BytesIO()
    with pd.ExcelWriter(excel_io, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name='JIRA Import', index=False)
        
        workbook = writer.book
        worksheet = writer.sheets['JIRA Import']
        
        header_format = workbook.add_format({
            'bold': True,
            'bg_color': '#D8E4BC',
            'border': 1
        })
        
        epic_format = workbook.add_format({
            'text_wrap': True,
            'valign': 'top',
            'bg_color': '#F0F3F7'  
        })
        
        story_format = workbook.add_format({
            'text_wrap': True,
            'valign': 'top',
            'bg_color': '#E8F3E8'  
        })
        
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)
        
        for row in range(len(df)):
            row_format = epic_format if is_epic.iloc[row] else story_format
            for col in range(len(df.columns)):
                worksheet.write(row + 1, col, df.iloc[row, col], row_format)
            
        worksheet.set_column('A:A', 40)  
        worksheet.set_column('B:B', 60)  
        worksheet.set_column('C:F', 15)  
        
        worksheet.set_default_row(100)
    
    excel_io.seek(0)
    return excel_io

@app.post("/generate-documents")
async def generate_documents(doc_request: DocumentRequest):
    """Generate requirement documents in Word and/or Excel format"""
    try:
        requirements = []
        epics = await get_jira_issues("issuetype = Epic")
        stories = await get_jira_issues("issuetype = Story")
        
        for item in epics + stories:
            req = {
                'id': item['key'],
                'summary': item['fields']['summary'],
                'description': item['fields'].get('description', ''),
                'type': 'FR' if item['fields']['issuetype']['name'] == 'Story' else 'NFR',
                'priority': item['fields'].get('priority', {}).get('name', 'Medium')
            }
            requirements.append(req)
        
        if doc_request.format == 'word':
            word_io = await generate_word_document(requirements)
            return StreamingResponse(
                word_io,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": "attachment; filename=requirements.docx"
                }
            )
            
        elif doc_request.format == 'excel':
            excel_io = await generate_excel_sheet(requirements)
            return StreamingResponse(
                excel_io,
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={
                    "Content-Disposition": "attachment; filename=jira_import.xlsx"
                }
            )
            
        else:  
            zip_io = BytesIO()
            with zipfile.ZipFile(zip_io, mode='w', compression=zipfile.ZIP_DEFLATED) as zip_file:
    
                word_io = await generate_word_document(requirements)
                zip_file.writestr('requirements.docx', word_io.getvalue())

                excel_io = await generate_excel_sheet(requirements)
                zip_file.writestr('jira_import.xlsx', excel_io.getvalue())
            
            zip_io.seek(0)
            return StreamingResponse(
                zip_io,
                media_type="application/zip",
                headers={
                    "Content-Disposition": "attachment; filename=requirement_documents.zip"
                }
            )
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate documents: {str(e)}"
        )

def parse_ai_response(raw_response: str) -> list:
    """Safely parse Gemini's JSON response with potential markdown formatting"""
    try:
        clean_response = re.sub(r'```json|```', '', raw_response)
        
        start = clean_response.find('[')
        end = clean_response.rfind(']') + 1
        
        if start == -1 or end == 0:
            raise ValueError("No JSON array found in response")
            
        json_str = clean_response[start:end]
        
        parsed = json.loads(json_str)
        
        if not isinstance(parsed, list):
            raise ValueError("Response is not a JSON array")
            
        return parsed
        
    except Exception as e:
        print("Failed to parse AI response:")
        print(f"Raw response: {raw_response}")
        print(f"Clean attempt: {json_str if 'json_str' in locals() else 'N/A'}")
        raise

def process_srs_file(file_path: str):
    """Process SRS file and create Jira stories"""
    try:
        # Read the file content
        with open(file_path, 'r') as f:
            file_content = f.read()
        
        # Process the requirements
        requirements = process_srs(file_content)
        
        print(f"Processing {len(requirements)} requirements...")
        
        # Create Jira issues for each requirement
        for req in requirements:
            try:
                # Map requirement type to issue type
                issue_type_map = {
                    "FR": "Story",
                    "NFR": "Task",
                    "SR": "Task"
                }
                
                # Extract type from ID (e.g., "FR-1" -> "FR")
                req_type = req["id"].split("-")[0]
                
                issue = JiraIssue(
                    summary=f"{req['id']}: {req['title']}",
                    description=(
                        f"*Requirement ID:* {req['id']}\n\n"
                        f"*Type:* {req_type}\n\n"
                        f"*Description:*\n{req['description']}\n\n"
                    ),
                    issue_type=issue_type_map.get(req_type, "Story"),
                    priority="Medium"
                )
                response = create_jira_issue(issue)
                print(f"Created Jira issue: {response['key']} - {req['title']}")
            except Exception as e:
                print(f"Error creating Jira issue for {req['id']}: {str(e)}")
                continue
        
        print("Completed processing requirements.")
        return True
    except Exception as e:
        print(f"Error processing SRS file: {str(e)}")
        raise

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
